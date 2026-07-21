"""Parse inbound customer questionnaires into a normalized list of questions.

Supports:
  - xlsx: uses openpyxl. Heuristically finds the header row and the
          "question" column; supports optional "section" / "control id"
          adjacent columns.
  - docx: extracts numbered paragraphs and questions ending in '?'.
  - pdf:  extracts text page by page, then applies the same heuristic.

Returns a list of ParsedItem dicts with `row_index` (stable position in the
source file), `question`, and optional `section` / `framework_ref`. The
row_index is what the export step uses to place the answer back into the
original file.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import List, Optional


# Regex heuristics for detecting question-like text.
_Q_HEADER_TOKENS = {
    "question", "questions", "control", "control question", "prompt",
    "assessment question", "requirement", "criteria",
}
_A_HEADER_TOKENS = {"answer", "response", "reply", "vendor response"}
_ID_HEADER_TOKENS = {"id", "control id", "ref", "reference", "#"}
_SECTION_HEADER_TOKENS = {"section", "domain", "category", "area", "topic"}

_QUESTION_PATTERN = re.compile(r"[?？]\s*$")
_NUMBERED_PATTERN = re.compile(r"^\s*(\d+(\.\d+)*\)?\s+|[a-z]\.\s+)", re.IGNORECASE)


@dataclass
class ParsedItem:
    row_index: int
    question: str
    section: Optional[str] = None
    framework_ref: Optional[str] = None


@dataclass
class ParseResult:
    file_kind: str                 # xlsx | docx | pdf
    items: List[ParsedItem]
    meta: dict                     # header row index, column map, etc.

    def as_dict(self) -> dict:
        return {
            "file_kind": self.file_kind,
            "items": [asdict(i) for i in self.items],
            "meta": self.meta,
        }


# ── public entry point ────────────────────────────────────────────────


def parse_questionnaire(path: str | Path) -> ParseResult:
    p = Path(path)
    ext = p.suffix.lower().lstrip(".")
    if ext in {"xlsx", "xlsm"}:
        return _parse_xlsx(p)
    if ext in {"docx"}:
        return _parse_docx(p)
    if ext in {"pdf"}:
        return _parse_pdf(p)
    raise ValueError(f"unsupported file type: .{ext}")


# ── xlsx ──────────────────────────────────────────────────────────────


def _parse_xlsx(p: Path) -> ParseResult:
    from openpyxl import load_workbook

    wb = load_workbook(filename=str(p), data_only=True, read_only=False)

    best_result: Optional[ParseResult] = None
    best_score = -1

    for sheet in wb.worksheets:
        # Scan first 20 rows for a header row containing a "question" token.
        header_row_idx = None
        column_map = {}
        for r in range(1, min(sheet.max_row, 20) + 1):
            row_vals = [
                str(sheet.cell(row=r, column=c).value or "").strip().lower()
                for c in range(1, sheet.max_column + 1)
            ]
            has_q = any(v in _Q_HEADER_TOKENS or "question" in v for v in row_vals)
            if has_q:
                header_row_idx = r
                for c, v in enumerate(row_vals, start=1):
                    if not v:
                        continue
                    if v in _Q_HEADER_TOKENS or "question" in v or "prompt" in v:
                        column_map.setdefault("question", c)
                    elif v in _A_HEADER_TOKENS or "answer" in v or "response" in v:
                        column_map.setdefault("answer", c)
                    elif v in _ID_HEADER_TOKENS or "control" in v and "id" in v:
                        column_map.setdefault("control_id", c)
                    elif v in _SECTION_HEADER_TOKENS:
                        column_map.setdefault("section", c)
                break

        if header_row_idx is None or "question" not in column_map:
            continue

        items: List[ParsedItem] = []
        q_col = column_map["question"]
        for r in range(header_row_idx + 1, sheet.max_row + 1):
            q_val = sheet.cell(row=r, column=q_col).value
            if q_val is None:
                continue
            q_str = str(q_val).strip()
            if len(q_str) < 5:
                continue

            section = None
            if "section" in column_map:
                s = sheet.cell(row=r, column=column_map["section"]).value
                if s is not None:
                    section = str(s).strip() or None
            framework_ref = None
            if "control_id" in column_map:
                c_id = sheet.cell(row=r, column=column_map["control_id"]).value
                if c_id is not None:
                    framework_ref = str(c_id).strip() or None

            items.append(
                ParsedItem(row_index=r, question=q_str, section=section, framework_ref=framework_ref)
            )

        if len(items) > best_score:
            best_score = len(items)
            best_result = ParseResult(
                file_kind="xlsx",
                items=items,
                meta={
                    "sheet_name": sheet.title,
                    "header_row": header_row_idx,
                    "column_map": column_map,
                    "answer_col": column_map.get("answer") or (q_col + 1),
                },
            )

    if not best_result or not best_result.items:
        raise ValueError("could not find a question column in any sheet")
    return best_result


# ── docx ──────────────────────────────────────────────────────────────


def _parse_docx(p: Path) -> ParseResult:
    from docx import Document

    doc = Document(str(p))
    items: List[ParsedItem] = []
    current_section: Optional[str] = None
    idx = 0

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue

        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            current_section = text
            continue

        if _looks_like_question(text):
            idx += 1
            items.append(ParsedItem(row_index=idx, question=text, section=current_section))

    # Also scan tables — many questionnaires ship as docx tables.
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if not text or not _looks_like_question(text):
                    continue
                idx += 1
                items.append(ParsedItem(row_index=idx, question=text, section=current_section))

    if not items:
        raise ValueError("no question-like text found in docx")
    return ParseResult(
        file_kind="docx",
        items=items,
        meta={"total_paragraphs": len(doc.paragraphs), "total_tables": len(doc.tables)},
    )


# ── pdf ──────────────────────────────────────────────────────────────


def _parse_pdf(p: Path) -> ParseResult:
    from pypdf import PdfReader

    reader = PdfReader(str(p))
    items: List[ParsedItem] = []
    idx = 0
    current_section: Optional[str] = None

    for page_no, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        for raw_line in text.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            # Very short lines that are all caps → treat as section headers.
            if len(line) < 60 and line.isupper():
                current_section = line.title()
                continue
            if _looks_like_question(line):
                idx += 1
                items.append(
                    ParsedItem(
                        row_index=idx,
                        question=line,
                        section=current_section,
                        framework_ref=f"p{page_no}",
                    )
                )

    if not items:
        raise ValueError("no question-like text found in pdf")
    return ParseResult(
        file_kind="pdf",
        items=items,
        meta={"pages": len(reader.pages)},
    )


# ── heuristics ────────────────────────────────────────────────────────


def _looks_like_question(text: str) -> bool:
    if len(text) < 10 or len(text) > 800:
        return False
    if _QUESTION_PATTERN.search(text):
        return True
    if _NUMBERED_PATTERN.match(text) and len(text) > 25:
        return True
    lowered = text.lower()
    starters = (
        "do you", "does the", "have you", "how do you", "how does",
        "is there", "are there", "please describe", "describe your",
        "provide", "list ", "explain ", "what ", "which ", "can you",
    )
    return any(lowered.startswith(s) for s in starters)
